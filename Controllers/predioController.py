import os
from fastapi import APIRouter, HTTPException, Depends, File, UploadFile, Form, BackgroundTasks
from sqlalchemy.orm import Session
from Models.predio import Predio, PredioDB, Predioform
from Models.rutas import RutasDB
from Models.adjuntos import AdjuntosDB
from Models.transaccion import TransaccionDB
from Models.propietarios import PropietariosDB
from database import SessionLocal
from docx import Document
from docx.shared import Pt  # Importar Pt para establecer tamaños en puntos
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Importar WD_ALIGN_PARAGRAPH para la alineación
from io import BytesIO
from typing import List, Annotated
from fastapi.responses import StreamingResponse, FileResponse
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

from fastapi.responses import Response
import copy
import shutil
from datetime import datetime
import locale
import fitz  # PyMuPDF
import asyncio
import tempfile
from docx.shared import Inches
import math
import json
from typing import Dict


# Registrar la fuente Arialic Hollow
pdfmetrics.registerFont(TTFont('Arialic Hollow', 'Fuentes/Arialic Hollow.ttf'))
router = APIRouter()

# Cargar la plantilla al inicio, fuera de la función del controlador
template_path = "existencia/noexiste.docx"
template_doc = Document(template_path)
# Establecer la configuración regional a español
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@router.get("/api/predios/", response_model=List[Predio], tags=["predios"])
def listar_predios(db: Session = Depends(get_db)):
    return db.query(PredioDB).all()

@router.get("/api/predios/{predio_id}", response_model=Predio, tags=["predios"])
def seleccionar_predio(predio_id: int, db: Session = Depends(get_db)):
    predio = db.query(PredioDB).filter(PredioDB.IdPredio == predio_id).first()
    if predio is None:
        raise HTTPException(status_code=404, detail="Predio no encontrado")
    return predio


# API para añadir predio
@router.post("/api/{id_u}/predios/", response_model=Predio, tags=["predios"])
async def añadir_predio(
        id_u: int,  # El ID del usuario que está realizando la acción
        predio: Annotated[Predioform, File()],
        db: Session = Depends(get_db)
):
    try:
        propietario = db.query(PropietariosDB).filter(PropietariosDB.Ci == predio.IdPropietario).first()
        if propietario is None:
            raise HTTPException(status_code=404, detail="Propietario no encontrado")
        predio.CodCatastral = (f"{predio.SubD}"
                               f"-{predio.Manzano}"
                               f"-{predio.NumeroPredio}"
                               f"-{predio.Uso}"
                               f"-{predio.Bloque}"
                               f"-{predio.Planta}"
                               f"-{predio.UnidadCat}")

        # Obtener el último valor de Ruta de la tabla Rutas
        ultima_ruta = db.query(RutasDB).order_by(RutasDB.Id.desc()).first()
        if ultima_ruta:
            predio.Ruta = f"{ultima_ruta.Direccion}/{predio.CodCatastral}"

        # Guarda el archivo en una ruta específica
        file_extension = predio.file.filename.split(".")[-1]
        path = f"{ultima_ruta.Direccion}/{predio.CodCatastral}"
        file_location = os.path.join(path, f"1.{file_extension}")
        # Crea el directorio si no existe
        os.makedirs(path, exist_ok=True)

        with open(file_location, "wb") as f:
            shutil.copyfileobj(predio.file.file, f)

        # Crea el objeto de predio en la base de datos
        db_predio = PredioDB(**predio.dict(exclude={"IdPredio", "file"}))
        db.add(db_predio)
        db.commit()
        db.refresh(db_predio)

        # Registrar la transacción de añadir predio
        db_transaccion = TransaccionDB(
            IDUsuario=id_u,  # El usuario que realiza la acción
            FechaHora=datetime.now(),
            TipoAccion="Añadir predio",  # Tipo de acción
            Descripcion=f"El usuario {id_u} añadió un nuevo predio con código catastral {predio.CodCatastral}.",
            # Descripción
            EntidadAfectada="Predio",  # La entidad afectada
            IDEntidadAfectada=db_predio.IdPredio  # El ID del nuevo predio
        )

        # Agregar la transacción
        db.add(db_transaccion)
        db.commit()  # Confirmar la transacción

        return db_predio

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al añadir el predio: {str(e)}")


# API para editar predio
@router.put("/api/{id_u}/predios/{predio_id}", response_model=Predio, tags=["predios"])
def editar_predio(id_u: int, predio_id: int, predio_actualizado: Predio, db: Session = Depends(get_db)):
    predio = db.query(PredioDB).filter(PredioDB.IdPredio == predio_id).first()
    if predio is None:
        raise HTTPException(status_code=404, detail="Predio no encontrado")

    # Actualizar el predio
    for key, value in predio_actualizado.dict(exclude={"IdPredio"}).items():
        setattr(predio, key, value)

    # Registrar la transacción de editar predio
    db_transaccion = TransaccionDB(
        IDUsuario=id_u,  # El usuario que realiza la acción
        FechaHora=datetime.now(),
        TipoAccion="Editar predio",  # Tipo de acción
        Descripcion=f"El usuario {id_u} editó el predio con ID {predio_id}.",  # Descripción
        EntidadAfectada="Predio",  # La entidad afectada
        IDEntidadAfectada=predio_id  # El ID del predio editado
    )

    # Agregar la transacción
    db.add(db_transaccion)
    db.commit()  # Confirmar la transacción

    db.commit()
    db.refresh(predio)

    return predio


# API para eliminar predio
@router.delete("/api/{id_u}/predios/{predio_id}", tags=["predios"])
def eliminar_predio(id_u: int, predio_id: int, db: Session = Depends(get_db)):
    predio = db.query(PredioDB).filter(PredioDB.IdPredio == predio_id).first()
    if predio is None:
        raise HTTPException(status_code=404, detail="Predio no encontrado")

    # Ruta del directorio que deseas eliminar
    path = predio.Ruta

    # Eliminar el directorio y todo su contenido
    try:
        shutil.rmtree(path)
        print(f"Directorio '{path}' eliminado con éxito.")
    except Exception as e:
        print(f"Error al eliminar el directorio: {str(e)}")

    # Registrar la transacción de eliminar predio
    db_transaccion = TransaccionDB(
        IDUsuario=id_u,  # El usuario que realiza la acción
        FechaHora=datetime.now(),
        TipoAccion="Eliminar predio",  # Tipo de acción
        Descripcion=f"El usuario {id_u} eliminó el predio con ID {predio_id}.",  # Descripción
        EntidadAfectada="Predio",  # La entidad afectada
        IDEntidadAfectada=predio_id  # El ID del predio eliminado
    )

    # Agregar la transacción
    db.add(db_transaccion)
    db.commit()  # Confirmar la transacción

    # Eliminar el predio de la base de datos
    db.delete(predio)

    # Eliminar todos los adjuntos relacionados al predio
    db.query(AdjuntosDB).filter(AdjuntosDB.IdPredio == predio.IdPredio).delete(synchronize_session=False)
    db.commit()

    return {"detail": "Predio eliminado", "predio": predio}

# Endpoint para combinar archivos en un solo PDF
@router.get("/api/{id_u}/predios/{id}/archivos/combinados", tags=["predios"])
async def combine_archivos_predio(
    id_u: int,  # ID del usuario que realiza la acción
    id: int,  # ID del predio
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    db_predio = db.query(PredioDB).filter(PredioDB.IdPredio == id).first()
    if db_predio is None:
        raise HTTPException(status_code=404, detail="Predio no encontrado")

    directory_path = db_predio.Ruta
    if not os.path.exists(directory_path):
        raise HTTPException(status_code=404, detail="La ruta del predio no existe.")

    archivos = [f for f in os.listdir(directory_path) if f.endswith((".pdf", ".jpg", ".jpeg", ".png", "bmp", "tif"))]
    if not archivos:
        raise HTTPException(status_code=404, detail="No se encontraron archivos en la ruta.")

    # Crear un archivo temporal
    temp_file_path = tempfile.mktemp(suffix='.pdf')
    pdf_writer = fitz.open()

    async def process_file(file):
        path = os.path.join(directory_path, file)
        if file.endswith(".pdf"):
            pdf_writer.insert_pdf(fitz.open(path))
        else:
            img_pdf = fitz.open()
            img_pdf.new_page()  # Crea una nueva página
            img_pdf[-1].insert_image(img_pdf[-1].rect, filename=path)  # Inserta la imagen
            pdf_writer.insert_pdf(img_pdf)  # Inserta el PDF creado con la imagen

    await asyncio.gather(*(process_file(f) for f in archivos))
    add_watermark_to_pdf(pdf_writer, "G.A.M COCHABAMBA")

    try:
        pdf_writer.save(temp_file_path)
    finally:
        pdf_writer.close()

    # Registrar la transacción de combinar archivos
    db_transaccion = TransaccionDB(
        IDUsuario=id_u,  # El usuario que realiza la acción
        FechaHora=datetime.now(),
        TipoAccion="Combinar archivos",  # Tipo de acción
        Descripcion=f"El usuario {id_u} combinó los archivos del predio con ID {id}.",  # Descripción
        EntidadAfectada="Predio",  # La entidad afectada
        IDEntidadAfectada=id  # El ID del predio
    )

    # Agregar la transacción
    db.add(db_transaccion)
    db.commit()  # Confirmar la transacción

    # Agregar la tarea de eliminación del archivo a BackgroundTasks
    background_tasks.add_task(remove_temp_file, temp_file_path)

    return FileResponse(temp_file_path, media_type="application/pdf", filename="archivos_combinados.pdf")

def add_watermark_to_pdf(pdf_writer, watermark_text):
    font_size = 25
    color = (0, 0, 0, 0.5)
    spacing_x = 200
    spacing_y = 100
    font_file1 = r"Fuentes/Arialic Hollow.ttf"

    for page in pdf_writer:
        width, height = page.rect.width, page.rect.height
        for x in range(0, int(width), spacing_x):
            for y in range(0, int(height), spacing_y):
                page.insert_font(fontfile=font_file1, fontname="F0")
                page.insert_text((x, y), watermark_text, fontsize=font_size, color=color, rotate=45, render_mode=0, fontname="F0")

def remove_temp_file(file_path):
    try:
        os.remove(file_path)
    except Exception as e:
        print(f"Error al eliminar el archivo temporal: {e}")


# Endpoint para generar el documento y registrar la transacción
@router.post("/api/{id_u}/predios/noexistencia", tags=["predios"])
async def generar_noexistencia(id_u: int, codCatastral: str, ci: int, db: Session = Depends(get_db)):
    partes = codCatastral.split('-')
    if len(partes) < 7:
        raise HTTPException(status_code=400, detail="El codCatastral no tiene el formato correcto.")

    try:
        # Crear una copia de la plantilla para cada solicitud
        doc = copy.deepcopy(template_doc)

        # Modificar el documento según tus necesidades
        fecha_actual = datetime.now().strftime("%d de %B de %Y")
        sustituciones = {
            "{subd}": partes[0],
            "{manzano}": partes[1],
            "{numpredio}": partes[2],
            "{date}": fecha_actual,
            "{ci}": str(ci),
        }

        for para in doc.paragraphs:
            for key, value in sustituciones.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)

        # Guardar el documento en un objeto BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Registrar la transacción
        db_transaccion = TransaccionDB(
            IDUsuario=id_u,  # El usuario que realizó la acción
            FechaHora=datetime.now(),
            TipoAccion="Generar documento no existencia",  # Acción de generar documento
            Descripcion=f"El usuario {id_u} generó el documento de no existencia para el predio con codCatastral {codCatastral}.",  # Descripción de la acción
            EntidadAfectada="Predio",  # La entidad afectada es un predio
            IDEntidadAfectada=partes[3]  # ID del predio basado en el código catastral
        )

        # Agregar la transacción
        db.add(db_transaccion)
        db.commit()  # Confirmar la transacción

        # Devolver el documento como respuesta
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=nuevo_documento.docx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo generar el documento: {str(e)}")


template_path2 = "existencia/existe2.docx"


# Endpoint para generar el documento de existencia y registrar la transacción
@router.post("/api/{id_u}/predios/existencia", tags=["predios"])
async def generar_existencia(id_u: int, codCatastral: str, db: Session = Depends(get_db)):
    predio = db.query(PredioDB).filter(PredioDB.CodCatastral == codCatastral).first()
    partes = codCatastral.split('-')
    if len(partes) < 7:
        raise HTTPException(status_code=400, detail="El codCatastral no tiene el formato correcto.")

    try:
        # Crear una copia del documento plantilla
        doc = Document(template_path2)

        # Obtener la fecha actual
        fecha_actual = datetime.now().strftime("%d de %B de %Y")
        """sustituciones = {
            "{subd}": partes[1],
            "{mza}": partes[2],
            "{idpredio}": partes[3],
            "{date}": fecha_actual,
        }"""

        sustituciones = {
            "{subd}": str(predio.SubD),
            "{mza}": str(predio.Manzano),
            "{numpred}": str(predio.NumeroPredio),
            "{fechaaprob}": str(predio.FechaAprobacion),
            "{fechadoc}": predio.FechaDocumentoAprobador,
            "{tipopredio}": predio.TipoPredio,
            "{ci}": str(predio.IdPropietario),
            "{superficie}": str(predio.Superficie),
            "{unidadcat}": str(predio.UnidadCat),
            "{date}": fecha_actual,
        }

        # Sustituir texto en los párrafos del documento
        for para in doc.paragraphs:
            for key, value in sustituciones.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)

        # Sustituir texto en las celdas de las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in sustituciones.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

                    # Ajustar el espaciado y formato después de la sustitución
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alineación izquierda

        # Ruta del archivo PDF que ya está en el servidor
        pdf_path = f"{predio.Ruta}/1.pdf"  # Reemplaza con la ruta donde se encuentra el PDF

        # Verificar si el archivo PDF existe
        if not os.path.exists(pdf_path):
            raise HTTPException(status_code=404,
                                detail="El archivo PDF no se encuentra en la ruta especificada.")

        # Abrir y procesar el archivo PDF desde la ruta
        doc_pdf = fitz.open(pdf_path)
        first_page = doc_pdf[0]  # Obtener la primera página
        pix = first_page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Escalado de 2x para mayor resolución
        img_bytes = pix.tobytes("png")  # Convertir la página a PNG en bytes

        # Guardar la imagen en un archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_image:
            temp_image.write(img_bytes)
            temp_image_path = temp_image.name

        # Asegurarse de cerrar el archivo PDF para evitar el bloqueo
        doc_pdf.close()

        # Insertar la imagen en el documento, en la primera tabla o al final
        tables = doc.tables
        if tables:
            table = tables[1]  # Elegimos la primera tabla (ajustar según sea necesario)
            cell = table.cell(0, 0)  # Suponiendo que el cuadro de texto está en la primera celda
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(temp_image_path, width=Inches(2.5))  # Ajusta el tamaño de la imagen según sea necesario

        # Guardar el documento editado en un objeto BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Eliminar los archivos temporales
        os.remove(temp_image_path)

        # Registrar la transacción
        db_transaccion = TransaccionDB(
            IDUsuario=id_u,  # El usuario que realizó la acción
            FechaHora=datetime.now(),
            TipoAccion="Generar documento de existencia",  # Acción de generar documento de existencia
            Descripcion=f"El usuario {id_u} generó el documento de existencia para el predio con codCatastral {codCatastral}.",  # Descripción de la acción
            EntidadAfectada="Predio",  # La entidad afectada es un predio
            IDEntidadAfectada=partes[3]  # ID del predio basado en el código catastral
        )

        # Agregar la transacción
        db.add(db_transaccion)
        db.commit()  # Confirmar la transacción

        # Devolver el documento como respuesta
        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=nuevo_documento.docx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo generar el documento: {str(e)}")